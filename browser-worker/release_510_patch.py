from __future__ import annotations

import re
from pathlib import Path

APP = Path('/app/app.py')
ADMIN = Path('/app/admin_dashboard.py')
MANIFEST = Path('/app/release_manifest.py')
MODULE = Path('/app/project_intelligence.py')
VERSION = '6.0.0'

MODULE_SOURCE = r'''from __future__ import annotations
import hashlib,json,re,uuid
from datetime import UTC,datetime
from pathlib import Path
from typing import Any
from docx import Document
from fastapi import APIRouter,File,Form,HTTPException,UploadFile
from fastapi.responses import HTMLResponse,PlainTextResponse
from openpyxl import load_workbook
from pydantic import BaseModel,Field
from pypdf import PdfReader
VERSION="6.0.0"
ROOT=Path("/app/artifacts/project-intelligence"); PROJECTS=ROOT/"projects"; UPLOADS=ROOT/"uploads"; VERSIONS=ROOT/"versions"
for x in (PROJECTS,UPLOADS,VERSIONS): x.mkdir(parents=True,exist_ok=True)
router=APIRouter(prefix="/project-intelligence",tags=["Project Intelligence"])
TYPES={"onec_report":"Отчёт 1С","onec_print_form":"Печатная форма 1С","onec_change":"Доработка 1С","integration":"Интеграция","dashboard":"Dashboard","business_process":"Бизнес-процесс","rest_api":"REST API","custom":"Другой проект"}
class Answers(BaseModel): answers:dict[str,Any]=Field(default_factory=dict)
class Decision(BaseModel): approved:bool; comment:str=""
def now(): return datetime.now(UTC).isoformat()
def pf(i):
 if not re.fullmatch(r"[a-f0-9]{32}",i): raise HTTPException(400,"Invalid project id")
 return PROJECTS/f"{i}.json"
def load(i):
 p=pf(i)
 if not p.exists(): raise HTTPException(404,"Project not found")
 return json.loads(p.read_text(encoding="utf-8"))
def save(p,event=None):
 p["updated_at"]=now()
 if event:p.setdefault("history",[]).append({"at":now(),"event":event,"revision":p.get("revision",1)})
 pf(p["id"]).write_text(json.dumps(p,ensure_ascii=False,indent=2),encoding="utf-8")
 (VERSIONS/f"{p['id']}-r{int(p.get('revision',1)):04d}.json").write_text(json.dumps(p,ensure_ascii=False,indent=2),encoding="utf-8")
 return p
def parse(path):
 s=path.suffix.lower()
 if s==".docx":
  d=Document(path); pars=[x.text.strip() for x in d.paragraphs if x.text.strip()]; tables=[[[c.text.strip() for c in r.cells] for r in t.rows[:200]] for t in d.tables[:30]]
  return {"format":"docx","text":"\n".join(pars)[:300000],"paragraphs":pars[:2000],"tables":tables}
 if s==".pdf":
  r=PdfReader(str(path)); pages=[]
  for n,p in enumerate(r.pages[:200]):
   try:t=p.extract_text() or ""
   except Exception as e:t=f"[Ошибка извлечения: {e}]"
   pages.append({"page":n+1,"text":t[:30000]})
  text="\n".join(x["text"] for x in pages)
  return {"format":"pdf","text":text[:300000],"pages":pages,"page_count":len(r.pages),"ocr_required":not bool(text.strip())}
 if s in {".xlsx",".xlsm"}:
  w=load_workbook(path,data_only=False); sheets=[]
  for sh in w.worksheets[:20]:
   rows=[]
   for row in sh.iter_rows(max_row=min(sh.max_row or 1,300),max_col=min(sh.max_column or 1,100)):
    vals=[c.value if isinstance(c.value,(str,int,float,bool)) or c.value is None else str(c.value) for c in row]
    if any(v not in (None,"") for v in vals):rows.append(vals)
   sheets.append({"name":sh.title,"headers":rows[0] if rows else [],"sample_rows":rows[1:11],"rows":rows[:300]})
  return {"format":"excel","sheets":sheets,"sheet_count":len(sheets)}
 if s in {".txt",".md",".csv"}:
  t=path.read_text(encoding="utf-8",errors="replace")[:300000]; return {"format":"text","text":t}
 return {"format":s.lstrip("."),"status":"stored_not_parsed"}
def text(m):
 a=m.get("analysis",{})
 if a.get("text"):return str(a["text"])
 return "\n".join(" | ".join(str(v or "") for v in row) for sh in a.get("sheets",[]) for row in sh.get("rows",[])[:100])
def requirements(ms):
 out=[]; seen=set()
 for m in ms:
  for z in re.findall(r"(?:необходимо|требуется|должен|должна|должны|нужно|предусмотреть|реализовать)\s+[^.!?\n]{8,250}",text(m),re.I):
   z=re.sub(r"\s+"," ",z).strip(); k=z.casefold()
   if k not in seen:seen.add(k);out.append({"id":f"REQ-{len(out)+1:03d}","text":z,"source_file":m.get("filename"),"confidence":0.78})
 return out[:500]
def conflicts(reqs,corpus):
 o=[]; low=corpus.casefold()
 for pair,title in [(("текущий остаток","остаток на дату"),"Не определён момент остатка"),(("плановая прибыль","фактическая прибыль"),"Не определён вид прибыли"),(("по заказам","по реализациям"),"Разные документы-основания")]:
  if all(x in low for x in pair):o.append({"id":f"CON-{len(o)+1:03d}","severity":"high","title":title,"evidence":list(pair),"resolution_required":True})
 return o
def questions(p):
 q=[("goal","Цель","Какое решение должно приниматься по результату?",1),("users","Пользователи","Кто будет пользоваться результатом?",1),("source_of_truth","Данные","Что сейчас считается источником правильных данных?",1),("period","Период","Какая дата основная и за какой период строится результат?",1),("filters","Отборы","Какие отборы и исключения нужны?",0),("calculations","Расчёты","Какие показатели и по каким формулам рассчитываются?",1),("acceptance","Приёмка","На каком примере доказать корректность?",1),("performance","Нагрузка","Какой объём и допустимое время выполнения?",0)]
 low="\n".join(text(x) for x in p.get("materials",[])).casefold()
 if any(x in low for x in ("прибыл","марж","себестоим")):q.append(("profit_logic","Расчёты","Прибыль плановая или фактическая и откуда берётся себестоимость?",1))
 if "остат" in low:q.append(("stock_logic","Расчёты","Остаток текущий или на дату, учитывать ли резерв?",1))
 for c in p.get("conflicts",[]):q.append((f"resolve_{c['id']}","Конфликт",f"Уточните: {c['title']}. Какой вариант правильный?",1))
 return [{"id":i,"section":s,"question":t,"required":bool(r)} for i,s,t,r in q]
def snapshot():
 cs=[Path("/app/artifacts/bitrix-onec-integration/latest.json"),Path("/app/artifacts/external-sources/latest.json")]+sorted(Path("/app/artifacts").glob("**/latest*.json"),key=lambda p:p.stat().st_mtime,reverse=True)[:20]
 for p in cs:
  try:d=json.loads(p.read_text(encoding="utf-8"))
  except:continue
  t=json.dumps(d,ensure_ascii=False).casefold()
  if "mcp_1c" in t or "onec_profile" in t:return {"source":str(p),"payload":d}
 return {"source":None,"payload":{}}
def assess(p):
 corpus="\n".join(text(x) for x in p.get("materials",[]))+json.dumps(p.get("answers",{}),ensure_ascii=False); low=corpus.casefold(); snap=snapshot(); st=json.dumps(snap["payload"],ensure_ascii=False).casefold(); cand=[]
 for token,name,kind in [("заказ","Документ.ЗаказКлиента","Документ"),("реализац","Документ.РеализацияТоваровУслуг","Документ"),("контрагент","Справочник.Контрагенты","Справочник"),("номенклатур","Справочник.Номенклатура","Справочник"),("остат","РегистрНакопления.ТоварыНаСкладах","Регистр"),("себестоим","РегистрНакопления.СебестоимостьТоваров","Регистр")]:
  if token in low:cand.append({"name":name,"type":kind,"confirmed_in_snapshot":name.split(".")[-1].casefold() in st})
 required=[x["id"] for x in p.get("questions",[]) if x.get("required")]; missing=[x for x in required if not str(p.get("answers",{}).get(x,"")).strip()]; blockers=[c for c in p.get("conflicts",[]) if c.get("resolution_required") and not p.get("answers",{}).get(f"resolve_{c['id']}")]
 warnings=[]
 if not snap["source"]:warnings.append({"severity":"medium","code":"MCP-001","title":"Нет актуального снимка 1С"})
 if not p.get("answers",{}).get("acceptance"):warnings.append({"severity":"high","code":"ACC-001","title":"Нет контрольного примера"})
 readiness=max(0,min(100,round(100*(len(required)-len(missing))/max(1,len(required))-len(blockers)*8)))
 return {"status":"ready_for_specification" if not missing and not blockers else "clarification_required","readiness":readiness,"missing_answers":missing,"unresolved_conflicts":blockers,"mcp_evidence":{"available":bool(snap["source"]),"source":snap["source"]},"candidate_objects":cand,"warnings":warnings}
def docs(p):
 a=p.get("analysis") or assess(p); ans=p.get("answers",{}); base={"project_id":p["id"],"revision":p.get("revision",1),"title":p["title"],"status":"approved" if p.get("approval",{}).get("approved") else "draft"}
 return {"business":{**base,"document":"Пользовательское описание","goal":ans.get("goal"),"users":ans.get("users"),"description":p.get("description")},"technical":{**base,"document":"Техническое задание","requirements":p.get("requirements",[]),"data_sources":a.get("candidate_objects",[]),"answers":ans,"risks":a.get("warnings",[])+p.get("conflicts",[])},"tests":{**base,"document":"Критерии приёмки","reference":ans.get("acceptance"),"tests":["Сверить контрольный пример","Проверить пустые значения","Проверить права","Проверить нагрузку","Проверить экспорт"]},"open_questions":{**base,"document":"Открытые вопросы","missing":a.get("missing_answers",[]),"conflicts":a.get("unresolved_conflicts",[])}}
def md(d):return "# "+d.get("document","Документ")+"\n\n```json\n"+json.dumps(d,ensure_ascii=False,indent=2)+"\n```\n"
@router.get("",response_class=HTMLResponse)
def page():return HTML
@router.get("/api/projects")
def listing():
 rows=[]
 for x in sorted(PROJECTS.glob("*.json"),key=lambda p:p.stat().st_mtime,reverse=True):
  try:p=json.loads(x.read_text(encoding="utf-8"));rows.append({k:p.get(k) for k in ("id","title","project_type","status","revision","updated_at")})
  except:pass
 return {"version":VERSION,"projects":rows}
@router.post("/api/projects")
async def create(title:str=Form(...),project_type:str=Form("onec_report"),description:str=Form(""),files:list[UploadFile]=File(default=[])):
 i=uuid.uuid4().hex; ms=[]; total=0
 for n,u in enumerate(files[:30]):
  if not u.filename:continue
  b=await u.read();total+=len(b)
  if len(b)>40*1024*1024 or total>150*1024*1024:raise HTTPException(413,"Upload limit exceeded")
  name=re.sub(r"[^A-Za-zА-Яа-яЁё0-9._-]","_",Path(u.filename).name);target=UPLOADS/f"{i}-{n:02d}-{name}";target.write_bytes(b);ms.append({"id":uuid.uuid4().hex,"filename":u.filename,"size":len(b),"sha256":hashlib.sha256(b).hexdigest(),"analysis":parse(target)})
 req=requirements(ms);corpus="\n".join(text(x) for x in ms);p={"id":i,"version":VERSION,"revision":1,"title":title,"project_type":project_type,"project_type_label":TYPES.get(project_type,project_type),"description":description,"status":"interview_required","created_at":now(),"materials":ms,"requirements":req,"conflicts":conflicts(req,corpus),"answers":{},"history":[]};p["questions"]=questions(p);return save(p,"project_created")
@router.post("/api/projects/{i}/materials")
async def add(i:str,files:list[UploadFile]=File(...)):
 p=load(i)
 for u in files[:30]:
  b=await u.read();name=re.sub(r"[^A-Za-zА-Яа-яЁё0-9._-]","_",Path(u.filename or "file").name);t=UPLOADS/f"{i}-{uuid.uuid4().hex[:8]}-{name}";t.write_bytes(b);p["materials"].append({"id":uuid.uuid4().hex,"filename":u.filename,"size":len(b),"sha256":hashlib.sha256(b).hexdigest(),"analysis":parse(t)})
 p["requirements"]=requirements(p["materials"]);p["conflicts"]=conflicts(p["requirements"],"\n".join(text(x) for x in p["materials"]));p["questions"]=questions(p);p["revision"]+=1;return save(p,"materials_added")
@router.get("/api/projects/{i}")
def get(i:str):return load(i)
@router.post("/api/projects/{i}/answers")
def answer(i:str,r:Answers):
 p=load(i);p["answers"]={**p.get("answers",{}),**r.answers};p["revision"]+=1;p["status"]="ready_for_analysis";return save(p,"answers_updated")
@router.post("/api/projects/{i}/analyze")
def analyze(i:str):
 p=load(i);p["analysis"]=assess(p);p["status"]="prototype_review" if p["analysis"]["status"]=="ready_for_specification" else "clarification_required";return save(p,"analysis_completed")
@router.post("/api/projects/{i}/approve")
def approve(i:str,r:Decision):
 p=load(i);p["approval"]={"approved":r.approved,"comment":r.comment,"at":now()};p["documents"]=docs(p);p["status"]="documents_ready" if r.approved else "revision_required";return save(p,"project_decision")
@router.get("/api/projects/{i}/documents")
def documents(i:str):
 p=load(i);return p.get("documents") or docs(p)
@router.get("/api/projects/{i}/documents/{name}.md",response_class=PlainTextResponse)
def document(i:str,name:str):
 d=documents(i)
 if name not in d:raise HTTPException(404,"Document not found")
 return md(d[name])
HTML=r'''<!doctype html><html lang="ru"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>AI Project Analyst</title><style>body{margin:0;background:#f5f6f8;color:#17191c;font:14px/1.45 Inter,Segoe UI,sans-serif}main{max-width:1500px;margin:auto;padding:24px}.grid{display:grid;grid-template-columns:340px 1fr;gap:12px}.p{background:#fff;border:1px solid #e3e6ea;border-radius:11px;margin-bottom:10px}.h{padding:12px 14px;border-bottom:1px solid #e3e6ea;font-weight:800}.b{padding:14px}.item{padding:10px;border:1px solid #e3e6ea;border-radius:8px;margin:7px 0;cursor:pointer}input,select,textarea{width:100%;padding:9px;margin:5px 0 10px;border:1px solid #e3e6ea;border-radius:7px;box-sizing:border-box}.btn{padding:9px 12px;border:1px solid #e3e6ea;border-radius:7px;background:#fff}.primary{background:#202124;color:#fff}.tabs{display:flex;gap:4px;border-bottom:1px solid #e3e6ea}.tab{padding:10px;border:0;background:none}.tab.on{border-bottom:2px solid #5e6ad2;font-weight:800}.v{display:none;padding:14px}.v.on{display:block}.muted{color:#70757d}pre{white-space:pre-wrap;background:#f8f9fa;padding:10px}@media(max-width:900px){.grid{grid-template-columns:1fr}}</style></head><body><main><div><small>AI-BIT Enterprise 6.0</small><h1>AI Project Analyst</h1><p class="muted">Word/PDF/Excel → требования → конфликты → интервью → MCP → документы</p></div><div class="grid"><aside><div class="p"><div class="h">Новый проект</div><div class="b"><form id="c"><input name="title" placeholder="Название" required><select name="project_type"><option value="onec_report">Отчёт 1С</option><option value="onec_change">Доработка 1С</option><option value="integration">Интеграция</option><option value="custom">Другой</option></select><textarea name="description" placeholder="Описание"></textarea><input type="file" name="files" multiple><button class="btn primary">Создать</button></form></div></div><div class="p"><div class="h">Проекты</div><div class="b" id="list"></div></div></aside><section class="p"><div class="h" id="title">Выберите проект</div><div class="tabs"><button class="tab on" data-v="o">Обзор</button><button class="tab" data-v="m">Материалы</button><button class="tab" data-v="r">Требования</button><button class="tab" data-v="q">Интервью</button><button class="tab" data-v="a">AI Review</button><button class="tab" data-v="d">Документы</button></div><div id="o" class="v on"></div><div id="m" class="v"></div><div id="r" class="v"></div><div id="q" class="v"></div><div id="a" class="v"></div><div id="d" class="v"></div></section></div></main><script>const $=s=>document.querySelector(s);let p=null;async function api(u,o){let r=await fetch(u,o);if(!r.ok)throw Error(await r.text());return r.json()}async function list(){let x=await api('/project-intelligence/api/projects');$('#list').innerHTML=x.projects.map(z=>`<div class=item data-id=${z.id}><b>${z.title}</b><div class=muted>${z.status} · r${z.revision}</div></div>`).join('');document.querySelectorAll('[data-id]').forEach(x=>x.onclick=()=>openP(x.dataset.id))}async function openP(i){p=await api('/project-intelligence/api/projects/'+i);render()}function render(){if(!p)return;$('#title').textContent=p.title;$('#o').innerHTML=`<p>Статус: <b>${p.status}</b></p><p>Материалов: ${p.materials.length}; требований: ${p.requirements.length}; конфликтов: ${p.conflicts.length}</p><button class="btn primary" onclick=analyze()>Запустить AI Review</button> <button class=btn onclick=approve()>Утвердить</button>`;$('#m').innerHTML=`<form id=add><input type=file name=files multiple><button class=btn>Добавить</button></form>`+p.materials.map(x=>`<div class=item><b>${x.filename}</b><div class=muted>${x.analysis.format}</div></div>`).join('');$('#add').onsubmit=add;$('#r').innerHTML=p.requirements.map(x=>`<div class=item><b>${x.id}</b> ${x.text}</div>`).join('')||'Нет извлечённых требований';$('#q').innerHTML=`<form id=ans>${p.questions.map(x=>`<b>${x.question}</b><textarea name=${x.id}>${p.answers[x.id]||''}</textarea>`).join('')}<button class="btn primary">Сохранить</button></form>`;$('#ans').onsubmit=answers;$('#a').innerHTML=p.analysis?`<h3>Готовность ${p.analysis.readiness}%</h3><pre>${JSON.stringify(p.analysis,null,2)}</pre>`:'Анализ не запускался';$('#d').innerHTML=p.documents?Object.keys(p.documents).map(x=>`<div class=item><a target=_blank href=/project-intelligence/api/projects/${p.id}/documents/${x}.md>${p.documents[x].document}</a></div>`).join(''):'Документы появятся после утверждения'}async function analyze(){await api(`/project-intelligence/api/projects/${p.id}/analyze`,{method:'POST'});openP(p.id)}async function approve(){await api(`/project-intelligence/api/projects/${p.id}/approve`,{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({approved:true,comment:''})});openP(p.id)}async function answers(e){e.preventDefault();let f=new FormData(e.target),a={};for(let [k,v] of f)a[k]=v;await api(`/project-intelligence/api/projects/${p.id}/answers`,{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({answers:a})});openP(p.id)}async function add(e){e.preventDefault();await api(`/project-intelligence/api/projects/${p.id}/materials`,{method:'POST',body:new FormData(e.target)});openP(p.id)}$('#c').onsubmit=async e=>{e.preventDefault();p=await api('/project-intelligence/api/projects',{method:'POST',body:new FormData(e.target)});list();render()};document.querySelectorAll('.tab').forEach(x=>x.onclick=()=>{document.querySelectorAll('.tab,.v').forEach(y=>y.classList.remove('on'));x.classList.add('on');$('#'+x.dataset.v).classList.add('on')});list()</script></body></html>'''
'''

def log(s,m): print(f'[{s}] {m}')
def insert_after(text,anchor,value):
 if value in text:return text,False
 if anchor not in text:return text,False
 return text.replace(anchor,anchor+value,1),True

def patch_app(text):
 imports=(('from pydantic_settings import BaseSettings, SettingsConfigDict\n','from onec_requirement_architect import router as onec_requirement_architect_router\n'),('from onec_requirement_architect import router as onec_requirement_architect_router\n','from project_intelligence import router as project_intelligence_router\n'))
 for anchor,line in imports:
  if line not in text:
   text,ok=insert_after(text,anchor,line)
   if not ok:text=line+text
 include1='app.include_router(onec_requirement_architect_router)';include2='app.include_router(project_intelligence_router)'
 for line in (include1,include2):
  if line not in text:
   m=re.search(r'(?m)^app\s*=\s*FastAPI\([^\n]*\)\s*$',text)
   if not m:raise RuntimeError('Cannot locate FastAPI application')
   text=text[:m.end()]+'\n'+line+text[m.end():]
 return text

def patch_admin(text):
 nav='<button data-key="projectIntelligence"><span class="icon">PI</span><span class="label">AI Project Analyst</span></button>'
 if 'data-key="projectIntelligence"' not in text:
  m=re.search(r'(<button\b[^>]*data-key=["\']onecRequirements["\'][^>]*>)|(<button\b[^>]*data-key=["\']system["\'][^>]*>)|(</nav>)',text,re.I)
  if m:text=text[:m.start()]+nav+text[m.start():]
 frame='<iframe class="frame" data-key="projectIntelligence" data-src="/project-intelligence"></iframe>'
 if 'data-key="projectIntelligence" data-src="/project-intelligence"' not in text:
  m=re.search(r'(<iframe\b[^>]*data-key=["\']onecRequirements["\'][^>]*></iframe>)|(<iframe\b[^>]*data-key=["\']system["\'][^>]*></iframe>)',text,re.I)
  if m:text=text[:m.start()]+frame+text[m.start():]
 meta="projectIntelligence:{title:'AI Project Analyst',subtitle:'Материалы, требования, интервью, MCP и документы',url:'/project-intelligence'},"
 if 'projectIntelligence:{title:' not in text:
  m=re.search(r'(onecRequirements:\{title:)|(system:\{title:)|(const meta=\{)',text)
  if m:text=text[:m.start()]+meta+text[m.start():]
 return text

def version(text):
 for old in ('5.1.0','5.0.0','4.7.0','4.6.0','4.5.0','4.4.0','4.3.3','3.2.1'):text=text.replace(old,VERSION)
 return text

def main():
 MODULE.write_text(MODULE_SOURCE,encoding='utf-8');compile(MODULE_SOURCE,str(MODULE),'exec');log('OK','Project Intelligence module generated')
 APP.write_text(version(patch_app(APP.read_text(encoding='utf-8'))),encoding='utf-8')
 ADMIN.write_text(version(patch_admin(ADMIN.read_text(encoding='utf-8'))),encoding='utf-8')
 if MANIFEST.exists():
  x=version(MANIFEST.read_text(encoding='utf-8'));x=re.sub(r'EDITION\s*=\s*"[^"]*"','EDITION = "Project Intelligence Platform"',x,count=1);MANIFEST.write_text(x,encoding='utf-8')
 body=APP.read_text(encoding='utf-8')
 for marker in ('onec_requirement_architect_router','project_intelligence_router','app.include_router(project_intelligence_router)',VERSION):
  if marker not in body:raise RuntimeError(f'6.0.0 missing critical marker: {marker}')
 log('OK','API registration validated')
 print('Applied AI-BIT Enterprise 6.0.0 — Project Intelligence Platform')
if __name__=='__main__':main()
