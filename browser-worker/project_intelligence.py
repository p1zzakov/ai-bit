from __future__ import annotations

import hashlib
import json
import re
import uuid
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document
from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from fastapi.responses import HTMLResponse, PlainTextResponse
from openpyxl import load_workbook
from pydantic import BaseModel, Field
from pypdf import PdfReader

VERSION = "6.0.0"
ROOT = Path("/app/artifacts/project-intelligence")
PROJECTS = ROOT / "projects"
UPLOADS = ROOT / "uploads"
VERSIONS = ROOT / "versions"
for folder in (PROJECTS, UPLOADS, VERSIONS):
    folder.mkdir(parents=True, exist_ok=True)

router = APIRouter(prefix="/project-intelligence", tags=["Project Intelligence"])

PROJECT_TYPES = {
    "onec_report": "Отчёт 1С",
    "onec_print_form": "Печатная форма 1С",
    "onec_change": "Доработка конфигурации 1С",
    "integration": "Интеграция",
    "dashboard": "Dashboard",
    "business_process": "Бизнес-процесс",
    "rest_api": "REST API",
    "processing": "Обработка",
    "custom": "Произвольный проект",
}

class AnswersRequest(BaseModel):
    answers: dict[str, Any] = Field(default_factory=dict)

class DecisionRequest(BaseModel):
    approved: bool
    comment: str = ""


def now_iso() -> str:
    return datetime.now(UTC).isoformat()


def project_file(project_id: str) -> Path:
    if not re.fullmatch(r"[a-f0-9]{32}", project_id):
        raise HTTPException(status_code=400, detail="Invalid project id")
    return PROJECTS / f"{project_id}.json"


def load_project(project_id: str) -> dict[str, Any]:
    path = project_file(project_id)
    if not path.exists():
        raise HTTPException(status_code=404, detail="Project not found")
    return json.loads(path.read_text(encoding="utf-8"))


def save_project(project: dict[str, Any], event: str | None = None) -> dict[str, Any]:
    project["updated_at"] = now_iso()
    if event:
        project.setdefault("history", []).append({"at": now_iso(), "event": event, "revision": project.get("revision", 1)})
    project_file(str(project["id"])).write_text(json.dumps(project, ensure_ascii=False, indent=2), encoding="utf-8")
    snapshot = VERSIONS / f"{project['id']}-r{int(project.get('revision', 1)):04d}.json"
    snapshot.write_text(json.dumps(project, ensure_ascii=False, indent=2), encoding="utf-8")
    return project


def safe_name(name: str) -> str:
    return re.sub(r"[^A-Za-zА-Яа-яЁё0-9._-]", "_", Path(name).name)[:180]


def parse_docx(path: Path) -> dict[str, Any]:
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = []
    for table in doc.tables[:30]:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows[:200]]
        tables.append(rows)
    return {"format": "docx", "paragraphs": paragraphs[:2000], "tables": tables, "text": "\n".join(paragraphs)[:200000]}


def parse_pdf(path: Path) -> dict[str, Any]:
    reader = PdfReader(str(path))
    pages = []
    for index, page in enumerate(reader.pages[:200]):
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            text = f"[Ошибка извлечения страницы: {exc}]"
        pages.append({"page": index + 1, "text": text[:30000]})
    joined = "\n".join(x["text"] for x in pages)
    return {"format": "pdf", "pages": pages, "page_count": len(reader.pages), "text": joined[:300000], "ocr_required": not bool(joined.strip())}


def parse_excel(path: Path) -> dict[str, Any]:
    wb = load_workbook(path, read_only=False, data_only=False)
    sheets = []
    for ws in wb.worksheets[:20]:
        rows, formulas = [], []
        for row in ws.iter_rows(min_row=1, max_row=min(ws.max_row or 1, 300), max_col=min(ws.max_column or 1, 100)):
            values = []
            for cell in row:
                value = cell.value
                if isinstance(value, str) and value.startswith("="):
                    formulas.append({"cell": cell.coordinate, "formula": value})
                values.append(value if isinstance(value, (str, int, float, bool)) or value is None else str(value))
            if any(v not in (None, "") for v in values):
                rows.append(values)
        sheets.append({"name": ws.title, "rows": rows[:300], "headers": rows[0] if rows else [], "sample_rows": rows[1:11], "formulas": formulas[:300]})
    return {"format": "excel", "sheets": sheets, "sheet_count": len(sheets)}


def parse_text(path: Path) -> dict[str, Any]:
    text = path.read_text(encoding="utf-8", errors="replace")[:300000]
    return {"format": "text", "text": text, "lines": [x.strip() for x in text.splitlines() if x.strip()][:3000]}


def parse_material(path: Path) -> dict[str, Any]:
    suffix = path.suffix.casefold()
    if suffix == ".docx":
        return parse_docx(path)
    if suffix == ".pdf":
        return parse_pdf(path)
    if suffix in {".xlsx", ".xlsm"}:
        return parse_excel(path)
    if suffix in {".txt", ".md", ".csv"}:
        return parse_text(path)
    return {"format": suffix.lstrip(".") or "unknown", "status": "stored_not_parsed", "note": "Файл сохранён, но автоматический разбор формата пока не поддерживается."}


def material_text(material: dict[str, Any]) -> str:
    parsed = material.get("analysis") or {}
    if parsed.get("text"):
        return str(parsed["text"])
    if parsed.get("format") == "excel":
        parts = []
        for sheet in parsed.get("sheets", []):
            parts.append(str(sheet.get("name", "")))
            for row in sheet.get("rows", [])[:100]:
                parts.append(" | ".join(str(x or "") for x in row))
        return "\n".join(parts)
    return json.dumps(parsed, ensure_ascii=False)


def extract_requirements(materials: list[dict[str, Any]]) -> list[dict[str, Any]]:
    result, seen = [], set()
    patterns = (
        r"(?:необходимо|требуется|должен|должна|должны|нужно|предусмотреть|реализовать)\s+[^.!?\n]{8,250}",
        r"(?:отч[её]т|форма|система|обработка|интеграция)\s+долж[^.!?\n]{8,250}",
    )
    for material in materials:
        text = material_text(material)
        for pattern in patterns:
            for match in re.findall(pattern, text, flags=re.I):
                normalized = re.sub(r"\s+", " ", match).strip()
                key = normalized.casefold()
                if key in seen:
                    continue
                seen.add(key)
                result.append({"id": f"REQ-{len(result)+1:03d}", "text": normalized, "source_file": material.get("filename"), "status": "extracted", "confidence": 0.78})
    return result[:500]


def detect_conflicts(requirements: list[dict[str, Any]], corpus: str) -> list[dict[str, Any]]:
    conflicts = []
    lower = corpus.casefold()
    pairs = [
        (("текущий остаток", "остаток на дату"), "Не определён момент расчёта остатка"),
        (("плановая прибыль", "фактическая прибыль"), "Не определён вид прибыли"),
        (("по заказам", "по реализациям"), "Разные документы-основания отчёта"),
        (("ежедневно", "в реальном времени"), "Противоречивая периодичность обновления"),
    ]
    for tokens, title in pairs:
        if all(token in lower for token in tokens):
            conflicts.append({"id": f"CON-{len(conflicts)+1:03d}", "severity": "high", "title": title, "evidence": list(tokens), "resolution_required": True})
    duplicate_groups: dict[str, list[str]] = {}
    for req in requirements:
        signature = re.sub(r"\W+", "", req["text"].casefold())[:80]
        duplicate_groups.setdefault(signature, []).append(req["id"])
    for ids in duplicate_groups.values():
        if len(ids) > 1:
            conflicts.append({"id": f"DUP-{len(conflicts)+1:03d}", "severity": "low", "title": "Дублирующееся требование", "requirements": ids, "resolution_required": False})
    return conflicts


def build_questions(project: dict[str, Any]) -> list[dict[str, Any]]:
    corpus = "\n".join(material_text(x) for x in project.get("materials", []))
    lower = corpus.casefold()
    questions = [
        {"id": "goal", "section": "Цель", "question": "Какое бизнес-решение должно приниматься по результату разработки?", "required": True},
        {"id": "users", "section": "Пользователи", "question": "Кто будет пользоваться результатом и какие права у этих пользователей?", "required": True},
        {"id": "source_of_truth", "section": "Данные", "question": "Какие документы или регистры сейчас считаются источником правильных данных?", "required": True},
        {"id": "period", "section": "Период", "question": "Какая дата является основной и за какой период строится результат?", "required": True},
        {"id": "filters", "section": "Отборы", "question": "Какие обязательные отборы, исключения и ограничения нужны?", "required": False},
        {"id": "calculations", "section": "Расчёты", "question": "Какие показатели рассчитываются и по каким формулам?", "required": True},
        {"id": "acceptance", "section": "Приёмка", "question": "На каком контрольном примере можно доказать корректность результата?", "required": True},
        {"id": "performance", "section": "Нагрузка", "question": "Какой ожидается объём данных и допустимое время формирования?", "required": False},
    ]
    if any(x in lower for x in ("прибыл", "марж", "себестоим")):
        questions.append({"id": "profit_logic", "section": "Расчёты", "question": "Прибыль плановая или фактическая? В какой момент и из какого регистра берётся себестоимость?", "required": True})
    if "остат" in lower:
        questions.append({"id": "stock_logic", "section": "Расчёты", "question": "Остаток текущий или на дату? Учитывать резерв, доступный остаток и ожидаемые поступления?", "required": True})
    for conflict in project.get("conflicts", []):
        questions.append({"id": f"resolve_{conflict['id']}", "section": "Конфликт", "question": f"Уточните конфликт: {conflict['title']}. Какой вариант считать правильным?", "required": True})
    return questions


def latest_mcp_snapshot() -> dict[str, Any]:
    candidates = [Path("/app/artifacts/bitrix-onec-integration/latest.json"), Path("/app/artifacts/external-sources/latest.json")]
    candidates += sorted(Path("/app/artifacts").glob("**/latest*.json"), key=lambda p: p.stat().st_mtime, reverse=True)[:30]
    for path in candidates:
        try:
            payload = json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            continue
        text = json.dumps(payload, ensure_ascii=False).casefold()
        if "onec_profile" in text or "mcp_1c" in text or "управлениепредприятием" in text:
            return {"source": str(path), "payload": payload}
    return {"source": None, "payload": {}}


def assess(project: dict[str, Any]) -> dict[str, Any]:
    corpus = "\n".join(material_text(x) for x in project.get("materials", [])) + "\n" + json.dumps(project.get("answers", {}), ensure_ascii=False)
    lower = corpus.casefold()
    snapshot = latest_mcp_snapshot()
    snapshot_text = json.dumps(snapshot.get("payload", {}), ensure_ascii=False).casefold()
    candidates = []
    mapping = {
        "заказ": ("Документ.ЗаказКлиента", "Документ"), "реализац": ("Документ.РеализацияТоваровУслуг", "Документ"),
        "контрагент": ("Справочник.Контрагенты", "Справочник"), "партнер": ("Справочник.Партнеры", "Справочник"),
        "номенклатур": ("Справочник.Номенклатура", "Справочник"), "остат": ("РегистрНакопления.ТоварыНаСкладах", "Регистр накопления"),
        "себестоим": ("РегистрНакопления.СебестоимостьТоваров", "Регистр накопления"), "оплат": ("Документ.БезналичноеПоступлениеДенежныхСредств", "Документ"),
    }
    for token, (name, kind) in mapping.items():
        if token in lower:
            short = name.split(".", 1)[-1].casefold()
            candidates.append({"name": name, "type": kind, "confirmed_in_snapshot": short in snapshot_text, "reason": f"Найден маркер «{token}»"})
    required_ids = [q["id"] for q in project.get("questions", []) if q.get("required")]
    missing = [qid for qid in required_ids if not str(project.get("answers", {}).get(qid, "")).strip()]
    blockers = [c for c in project.get("conflicts", []) if c.get("resolution_required") and not str(project.get("answers", {}).get(f"resolve_{c['id']}", "")).strip()]
    warnings = []
    if not snapshot.get("source"):
        warnings.append({"severity": "medium", "code": "MCP-001", "title": "Не найден актуальный evidence-снимок 1С", "action": "Запустить MCP-сбор и повторить анализ."})
    if not project.get("answers", {}).get("acceptance"):
        warnings.append({"severity": "high", "code": "ACC-001", "title": "Нет контрольного примера", "action": "Предоставить эталонный период и ожидаемый результат."})
    readiness = max(0, min(100, round(100 * (len(required_ids) - len(missing)) / max(1, len(required_ids)) - len(blockers) * 8)))
    status = "ready_for_specification" if not missing and not blockers else "clarification_required"
    prototype_columns = []
    for material in project.get("materials", []):
        for sheet in (material.get("analysis", {}).get("sheets") or [])[:1]:
            for header in sheet.get("headers", [])[:80]:
                if header not in (None, ""):
                    prototype_columns.append({"name": str(header), "source": "Требует подтверждения", "calculation": "Не определён"})
    return {"status": status, "readiness": readiness, "missing_answers": missing, "unresolved_conflicts": blockers, "mcp_evidence": {"available": bool(snapshot.get("source")), "source": snapshot.get("source")}, "candidate_objects": candidates, "warnings": warnings, "prototype": {"columns": prototype_columns, "status": "awaiting_approval"}}


def build_documents(project: dict[str, Any]) -> dict[str, Any]:
    analysis = project.get("analysis") or assess(project)
    answers = project.get("answers", {})
    base = {"project_id": project["id"], "revision": project.get("revision", 1), "title": project["title"], "project_type": project.get("project_type"), "status": "approved" if project.get("approval", {}).get("approved") else "draft"}
    technical = {**base, "document": "Техническое задание", "goal": answers.get("goal"), "users": answers.get("users"), "functional_requirements": project.get("requirements", []), "data_sources": analysis.get("candidate_objects", []), "filters": answers.get("filters"), "calculations": answers.get("calculations"), "period": answers.get("period"), "performance": answers.get("performance"), "constraints": ["Не изменять типовую конфигурацию без отдельного согласования", "Не выполнять запись данных из аналитического модуля", "Неподтверждённый источник не считать фактом"], "risks": analysis.get("warnings", []) + project.get("conflicts", [])}
    business = {**base, "document": "Пользовательское описание", "goal": answers.get("goal"), "users": answers.get("users"), "expected_result": project.get("description"), "prototype": analysis.get("prototype")}
    tests = {**base, "document": "Критерии приёмки и тест-кейсы", "reference_example": answers.get("acceptance"), "tests": ["Сверить результат с согласованным контрольным примером", "Проверить пустые значения и отсутствие данных", "Проверить права разных ролей", "Проверить рабочий объём и время выполнения", "Проверить экспорт без потери итогов и типов"]}
    open_questions = {**base, "document": "Открытые вопросы", "missing_answers": analysis.get("missing_answers", []), "conflicts": analysis.get("unresolved_conflicts", []), "warnings": analysis.get("warnings", [])}
    return {"business": business, "technical": technical, "tests": tests, "open_questions": open_questions}


def markdown_document(doc: dict[str, Any]) -> str:
    lines = [f"# {doc.get('document', 'Документ')}", "", f"**Проект:** {doc.get('title')}", f"**Версия:** {doc.get('revision')}", f"**Статус:** {doc.get('status')}", ""]
    for key, value in doc.items():
        if key in {"document", "title", "revision", "status", "project_id"}:
            continue
        lines += [f"## {key.replace('_', ' ').title()}", "```json", json.dumps(value, ensure_ascii=False, indent=2), "```", ""]
    return "\n".join(lines)

@router.get("", response_class=HTMLResponse)
def page() -> str:
    return PROJECT_INTELLIGENCE_HTML

@router.get("/api/projects")
def list_projects() -> dict[str, Any]:
    rows = []
    for path in sorted(PROJECTS.glob("*.json"), key=lambda p: p.stat().st_mtime, reverse=True):
        try:
            p = json.loads(path.read_text(encoding="utf-8"))
            rows.append({k: p.get(k) for k in ("id", "title", "project_type", "status", "revision", "updated_at")})
        except Exception:
            pass
    return {"version": VERSION, "projects": rows}

@router.post("/api/projects")
async def create_project(title: str = Form(...), project_type: str = Form("onec_report"), description: str = Form(""), files: list[UploadFile] = File(default=[])) -> dict[str, Any]:
    if project_type not in PROJECT_TYPES:
        raise HTTPException(status_code=400, detail="Unknown project type")
    project_id = uuid.uuid4().hex
    materials = []
    total_size = 0
    for index, upload in enumerate(files[:30]):
        if not upload.filename:
            continue
        content = await upload.read()
        total_size += len(content)
        if len(content) > 40 * 1024 * 1024 or total_size > 150 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="File or project upload limit exceeded")
        name = safe_name(upload.filename)
        target = UPLOADS / f"{project_id}-{index:02d}-{name}"
        target.write_bytes(content)
        analysis = parse_material(target)
        materials.append({"id": uuid.uuid4().hex, "filename": upload.filename, "stored_as": str(target), "size": len(content), "content_type": upload.content_type, "sha256": hashlib.sha256(content).hexdigest(), "analysis": analysis})
    requirements = extract_requirements(materials)
    corpus = "\n".join(material_text(x) for x in materials)
    project = {"id": project_id, "version": VERSION, "revision": 1, "title": title.strip(), "project_type": project_type, "project_type_label": PROJECT_TYPES[project_type], "description": description.strip(), "status": "interview_required", "created_at": now_iso(), "updated_at": now_iso(), "materials": materials, "requirements": requirements, "conflicts": detect_conflicts(requirements, corpus), "answers": {}, "history": []}
    project["questions"] = build_questions(project)
    return save_project(project, "project_created")

@router.post("/api/projects/{project_id}/materials")
async def add_materials(project_id: str, files: list[UploadFile] = File(...)) -> dict[str, Any]:
    project = load_project(project_id)
    for index, upload in enumerate(files[:30]):
        content = await upload.read()
        name = safe_name(upload.filename or f"file-{index}")
        target = UPLOADS / f"{project_id}-{uuid.uuid4().hex[:8]}-{name}"
        target.write_bytes(content)
        project.setdefault("materials", []).append({"id": uuid.uuid4().hex, "filename": upload.filename, "stored_as": str(target), "size": len(content), "content_type": upload.content_type, "sha256": hashlib.sha256(content).hexdigest(), "analysis": parse_material(target)})
    project["requirements"] = extract_requirements(project["materials"])
    corpus = "\n".join(material_text(x) for x in project["materials"])
    project["conflicts"] = detect_conflicts(project["requirements"], corpus)
    project["questions"] = build_questions(project)
    project["revision"] = int(project.get("revision", 1)) + 1
    return save_project(project, "materials_added")

@router.get("/api/projects/{project_id}")
def get_project(project_id: str) -> dict[str, Any]:
    return load_project(project_id)

@router.post("/api/projects/{project_id}/answers")
def submit_answers(project_id: str, request: AnswersRequest) -> dict[str, Any]:
    project = load_project(project_id)
    project["answers"] = {**project.get("answers", {}), **request.answers}
    project["revision"] = int(project.get("revision", 1)) + 1
    project["status"] = "ready_for_analysis"
    return save_project(project, "answers_updated")

@router.post("/api/projects/{project_id}/analyze")
def analyze(project_id: str) -> dict[str, Any]:
    project = load_project(project_id)
    project["analysis"] = assess(project)
    project["status"] = "prototype_review" if project["analysis"]["status"] == "ready_for_specification" else "clarification_required"
    return save_project(project, "analysis_completed")

@router.post("/api/projects/{project_id}/approve")
def approve(project_id: str, request: DecisionRequest) -> dict[str, Any]:
    project = load_project(project_id)
    project["approval"] = {"approved": request.approved, "comment": request.comment, "at": now_iso()}
    project["documents"] = build_documents(project)
    project["status"] = "documents_ready" if request.approved else "revision_required"
    project["revision"] = int(project.get("revision", 1)) + (0 if request.approved else 1)
    return save_project(project, "project_decision")

@router.get("/api/projects/{project_id}/documents")
def documents(project_id: str) -> dict[str, Any]:
    project = load_project(project_id)
    return project.get("documents") or build_documents(project)

@router.get("/api/projects/{project_id}/documents/{document_name}.md", response_class=PlainTextResponse)
def document_markdown(project_id: str, document_name: str) -> str:
    docs = documents(project_id)
    if document_name not in docs:
        raise HTTPException(status_code=404, detail="Document not found")
    return markdown_document(docs[document_name])

PROJECT_INTELLIGENCE_HTML = r'''<!doctype html><html lang="ru"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>AI Project Analyst — AI-BIT 6.0</title><style>
:root{--bg:#f5f6f8;--s:#fff;--line:#e3e6ea;--text:#17191c;--muted:#70757d;--a:#5e6ad2;--ok:#16835d;--warn:#a9670b}*{box-sizing:border-box}body{margin:0;background:var(--bg);color:var(--text);font:14px/1.45 Inter,Segoe UI,sans-serif}main{max-width:1600px;margin:auto;padding:24px}.head{display:flex;justify-content:space-between;gap:20px;margin-bottom:16px}.ey{font-size:11px;font-weight:800;color:var(--a);letter-spacing:.12em;text-transform:uppercase}h1{margin:5px 0}.grid{display:grid;grid-template-columns:350px 1fr;gap:12px}.panel{background:var(--s);border:1px solid var(--line);border-radius:11px;overflow:hidden;margin-bottom:10px}.ph{padding:12px 14px;background:#fafbfc;border-bottom:1px solid var(--line);font-weight:800}.body{padding:14px}.item{padding:10px;border:1px solid var(--line);border-radius:8px;margin-bottom:7px;cursor:pointer}.item:hover,.item.active{border-color:var(--a);background:#f8f8ff}.badge{display:inline-block;padding:3px 7px;border-radius:999px;background:#eef0f2;font-size:10px;font-weight:800}.btn{border:1px solid var(--line);background:#fff;padding:8px 11px;border-radius:7px;cursor:pointer}.btn.primary{background:#202124;color:#fff}.btn.accent{background:var(--a);color:#fff;border-color:var(--a)}input,select,textarea{width:100%;padding:9px;border:1px solid var(--line);border-radius:7px;margin:5px 0 10px;font:inherit}textarea{min-height:80px}.tabs{display:flex;gap:4px;border-bottom:1px solid var(--line);overflow:auto}.tab{padding:9px 10px;border:0;background:none;cursor:pointer}.tab.active{border-bottom:2px solid var(--a);font-weight:800}.view{display:none;padding:14px}.view.active{display:block}.row{display:flex;justify-content:space-between;gap:12px;padding:7px 0;border-bottom:1px solid #eef0f2}.muted{color:var(--muted)}pre{white-space:pre-wrap;word-break:break-word;background:#f8f9fa;padding:10px;border-radius:8px}.actions{display:flex;gap:8px;flex-wrap:wrap}@media(max-width:900px){.grid{grid-template-columns:1fr}}</style></head><body><main><div class="head"><div><div class="ey">AI-BIT Enterprise 6.0</div><h1>AI Project Analyst</h1><div class="muted">Материалы → требования → конфликты → интервью → MCP → прототип → документы</div></div></div><div class="grid"><aside><div class="panel"><div class="ph">Новый проект</div><div class="body"><form id="create"><input name="title" placeholder="Название проекта" required><select name="project_type"><option value="onec_report">Отчёт 1С</option><option value="onec_print_form">Печатная форма</option><option value="onec_change">Доработка 1С</option><option value="integration">Интеграция</option><option value="dashboard">Dashboard</option><option value="business_process">Бизнес-процесс</option><option value="rest_api">REST API</option><option value="custom">Другой проект</option></select><textarea name="description" placeholder="Что нужно получить"></textarea><input type="file" name="files" multiple accept=".docx,.pdf,.xlsx,.xlsm,.txt,.md,.csv"><button class="btn primary">Создать проект</button></form></div></div><div class="panel"><div class="ph">Проекты</div><div class="body" id="projects"></div></div></aside><section class="panel"><div class="ph" id="projectTitle">Выберите проект</div><div class="tabs"><button class="tab active" data-v="overview">Обзор</button><button class="tab" data-v="materials">Материалы</button><button class="tab" data-v="requirements">Требования</button><button class="tab" data-v="interview">Интервью</button><button class="tab" data-v="analysis">AI Review</button><button class="tab" data-v="documents">Документы</button></div><div id="overview" class="view active"></div><div id="materials" class="view"></div><div id="requirements" class="view"></div><div id="interview" class="view"></div><div id="analysis" class="view"></div><div id="documents" class="view"></div></section></div></main><script>
const $=s=>document.querySelector(s), esc=v=>String(v??'').replace(/[&<>"']/g,c=>({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'}[c]));let current=null;async function api(u,o){const r=await fetch(u,o);if(!r.ok)throw new Error(await r.text());return r.json()}async function list(){const d=await api('/project-intelligence/api/projects');$('#projects').innerHTML=d.projects.map(p=>`<div class="item" data-id="${p.id}"><b>${esc(p.title)}</b><div class="muted">${esc(p.project_type)} · r${p.revision}</div><span class="badge">${esc(p.status)}</span></div>`).join('')||'<div class="muted">Проектов нет</div>';document.querySelectorAll('[data-id]').forEach(x=>x.onclick=()=>openP(x.dataset.id))}async function openP(id){current=await api('/project-intelligence/api/projects/'+id);render()}function render(){const p=current;if(!p)return;$('#projectTitle').textContent=p.title;$('#overview').innerHTML=`<div class="row"><span>Тип</span><b>${esc(p.project_type_label)}</b></div><div class="row"><span>Статус</span><b>${esc(p.status)}</b></div><div class="row"><span>Ревизия</span><b>${p.revision}</b></div><div class="row"><span>Материалов</span><b>${p.materials.length}</b></div><div class="row"><span>Требований</span><b>${p.requirements.length}</b></div><div class="row"><span>Конфликтов</span><b>${p.conflicts.length}</b></div><div class="actions" style="margin-top:12px"><button class="btn accent" onclick="analyze()">Запустить AI Review</button><button class="btn" onclick="approve(true)">Утвердить</button></div>`;$('#materials').innerHTML=`<form id="add"><input type="file" name="files" multiple><button class="btn">Добавить материалы</button></form>`+p.materials.map(x=>`<div class="item"><b>${esc(x.filename)}</b><div class="muted">${esc(x.analysis.format)} · ${x.size} bytes</div></div>`).join('');$('#add').onsubmit=addFiles;$('#requirements').innerHTML=p.requirements.map(x=>`<div class="item"><b>${x.id}</b> ${esc(x.text)}<div class="muted">${esc(x.source_file)}</div></div>`).join('')||'<div class="muted">Явные требования не извлечены.</div>';$('#interview').innerHTML=`<form id="answers">${p.questions.map(q=>`<label><b>${esc(q.section)}: ${esc(q.question)}</b><textarea name="${q.id}">${esc(p.answers[q.id]||'')}</textarea></label>`).join('')}<button class="btn primary">Сохранить ответы</button></form>`;$('#answers').onsubmit=saveAnswers;$('#analysis').innerHTML=p.analysis?`<div class="row"><span>Готовность</span><b>${p.analysis.readiness}%</b></div><div class="row"><span>Статус</span><b>${esc(p.analysis.status)}</b></div><h3>Объекты 1С</h3><pre>${esc(JSON.stringify(p.analysis.candidate_objects,null,2))}</pre><h3>Предупреждения</h3><pre>${esc(JSON.stringify(p.analysis.warnings,null,2))}</pre>`:'<div class="muted">Анализ ещё не запускался.</div>';$('#documents').innerHTML=p.documents?Object.keys(p.documents).map(k=>`<div class="item"><b>${esc(p.documents[k].document)}</b><div><a href="/project-intelligence/api/projects/${p.id}/documents/${k}.md" target="_blank">Открыть Markdown</a></div></div>`).join(''):'<div class="muted">Документы появятся после утверждения.</div>'}async function analyze(){await api(`/project-intelligence/api/projects/${current.id}/analyze`,{method:'POST'});await openP(current.id)}async function approve(ok){await api(`/project-intelligence/api/projects/${current.id}/approve`,{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({approved:ok,comment:''})});await openP(current.id)}async function saveAnswers(e){e.preventDefault();const fd=new FormData(e.target),a={};for(const [k,v] of fd)a[k]=v;await api(`/project-intelligence/api/projects/${current.id}/answers`,{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({answers:a})});await openP(current.id)}async function addFiles(e){e.preventDefault();await api(`/project-intelligence/api/projects/${current.id}/materials`,{method:'POST',body:new FormData(e.target)});await openP(current.id)}$('#create').onsubmit=async e=>{e.preventDefault();current=await api('/project-intelligence/api/projects',{method:'POST',body:new FormData(e.target)});await list();render()};document.querySelectorAll('.tab').forEach(b=>b.onclick=()=>{document.querySelectorAll('.tab,.view').forEach(x=>x.classList.remove('active'));b.classList.add('active');$('#'+b.dataset.v).classList.add('active')});list();
</script></body></html>'''
